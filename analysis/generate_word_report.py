#!/usr/bin/env python3
"""Generate Word document report per NEW_ASK.md requirements."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
from pathlib import Path
from datetime import datetime
from PIL import Image

def set_cell_background(cell, fill):
    """Set background color of a table cell."""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_word_report():
    """Create comprehensive Word document report."""
    
    print("=== GENERATING WORD DOCUMENT REPORT ===\n")
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "ACME Cultural Funding Analysis Report"
    doc.core_properties.author = "ACME Analysis Team"
    doc.core_properties.subject = "Analysis of Austin's Cultural Funding Landscape"
    
    # Title Page
    title = doc.add_heading('ACME Cultural Funding Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Analysis Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    
    date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(14)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    exec_summary = doc.add_paragraph(
        "This report presents a comprehensive analysis of Austin's cultural funding landscape based on "
        "community feedback collected through surveys, focus groups, and stakeholder interviews. "
        "The analysis reveals critical insights about funding needs, barriers, and opportunities "
        "for improving support to Austin's creative community."
    )
    
    # Load all data
    results_dir = Path("data/results")
    funding_dir = results_dir / "cultural_funding_analysis"
    
    # WHO Analysis Section
    doc.add_heading('WHO: Quantitative Analysis', 1)
    
    # Load WHO metrics
    with open(results_dir / "who_metrics.json", 'r') as f:
        who_data = json.load(f)
    
    # Add data collection summary
    doc.add_heading('Data Collection Methods', 2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Count'
    
    # Data rows
    methods = [
        ('Survey Responses', who_data['survey_responses']),
        ('Focus Groups', who_data['focus_groups']),
        ('1:1 Interviews', who_data['one_on_ones']),
        ('Listening Sessions', who_data['listening_sessions'])
    ]
    
    for i, (method, count) in enumerate(methods, 1):
        cells = table.rows[i].cells
        cells[0].text = method
        cells[1].text = str(count) if isinstance(count, int) else count
    
    # Share of Voice
    doc.add_heading('Share of Voice', 2)
    
    voice_para = doc.add_paragraph()
    voice_para.add_run('How many creatives: ').bold = True
    voice_para.add_run(f"{who_data['share_of_voice']['creatives']['count']:,} ({who_data['share_of_voice']['creatives']['percentage']}%)\n")
    voice_para.add_run('How many organizational staff: ').bold = True
    voice_para.add_run(f"{who_data['share_of_voice']['organizational_staff']['count']:,} ({who_data['share_of_voice']['organizational_staff']['percentage']}%)\n")
    voice_para.add_run('How many community members/patrons: ').bold = True
    voice_para.add_run(f"{who_data['share_of_voice']['community_patrons']['count']:,} ({who_data['share_of_voice']['community_patrons']['percentage']}%)")
    
    doc.add_page_break()
    
    # WHAT Analysis Section
    doc.add_heading('WHAT: Thematic Analysis', 1)
    
    # Load cultural funding themes
    with open(funding_dir / "cultural_funding_themes.json", 'r') as f:
        themes_data = json.load(f)
    
    doc.add_heading('Top 10 Themes on Cultural Funding', 2)
    
    # Add pie chart image
    pie_chart_path = funding_dir / "cultural_funding_themes_pie_chart.png"
    if pie_chart_path.exists():
        doc.add_picture(str(pie_chart_path), width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add themes with supporting sentences
    for theme_info in themes_data['top_10_cultural_funding_themes']:
        doc.add_heading(f"{theme_info['rank']}. {theme_info['theme']}", 3)
        
        freq_para = doc.add_paragraph()
        freq_para.add_run(f"Frequency: {theme_info['count']} mentions ({theme_info['percentage']}%)").italic = True
        
        doc.add_paragraph('Supporting evidence:')
        
        for i, sentence in enumerate(theme_info['supporting_sentences'], 1):
            if sentence:
                doc.add_paragraph(f"{i}. {sentence}", style='List Number')
    
    doc.add_page_break()
    
    # Program-specific themes
    doc.add_heading('Top 3 Themes per Program', 2)
    
    # Load program themes
    with open(funding_dir / "program_themes_analysis.json", 'r') as f:
        program_data = json.load(f)
    
    for program_info in program_data['programs_analyzed']:
        doc.add_heading(program_info['program'], 3)
        
        if program_info['top_3_themes']:
            for theme in program_info['top_3_themes']:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(f"{theme['theme']}").bold = True
                bullet.add_run(f" - {theme['count']} mentions ({theme['percentage']}%)")
        else:
            doc.add_paragraph("No specific recommendation data available for this program", style='List Bullet')
    
    doc.add_page_break()
    
    # Transportation - Parking Lot
    doc.add_heading('Transportation Comments - Parking Lot Item', 1)
    
    # Load transportation data
    with open(funding_dir / "transportation_parking_lot.json", 'r') as f:
        transport_data = json.load(f)
    
    transport_para = doc.add_paragraph(
        f"Transportation-related comments have been identified and separated for later analysis. "
        f"A total of {transport_data['total_transportation_comments']} comments "
        f"({transport_data['percentage_of_total_responses']}% of all responses) mentioned transportation barriers."
    )
    
    doc.add_heading('Category Breakdown', 3)
    
    # Create transportation table
    transport_table = doc.add_table(rows=len(transport_data['category_breakdown']) + 1, cols=2)
    transport_table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = transport_table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Count'
    
    # Data
    for i, (category, count) in enumerate(transport_data['category_breakdown'].items(), 1):
        cells = transport_table.rows[i].cells
        cells[0].text = category
        cells[1].text = str(count)
    
    doc.add_heading('Key Insights', 3)
    for insight in transport_data['key_insights']:
        doc.add_paragraph(insight, style='List Bullet')
    
    doc.add_paragraph(
        '\nRecommendation: Transportation and parking issues represent a significant barrier to arts '
        'and cultural participation in Austin. This topic warrants a separate, focused analysis.'
    )
    
    # Save document
    output_path = results_dir / "reports" / f"ACME_Cultural_Funding_Report_{datetime.now().strftime('%Y%m%d')}.docx"
    output_path.parent.mkdir(exist_ok=True, parents=True)
    doc.save(output_path)
    
    # Also save as LATEST
    latest_path = results_dir / "reports" / "ACME_Cultural_Funding_Report_LATEST.docx"
    doc.save(latest_path)
    
    print(f"✓ Word document report saved to: {output_path}")
    print(f"✓ Also saved as: {latest_path}")

if __name__ == "__main__":
    create_word_report()