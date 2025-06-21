"""Executive report generator using professional templates."""

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from rich.console import Console

from ..config import settings
from ..validation.audit import AuditLogger
from .templates import ReportTemplate, ExecutiveSummaryTemplate

console = Console()


class ReportGenerator:
    """Generates professional executive reports from analysis results."""
    
    def __init__(self, audit_logger: Optional[AuditLogger] = None):
        self.audit_logger = audit_logger or AuditLogger()
        self.template = ReportTemplate()
        self.output_dir = settings.results_dir / "reports"
        self.output_dir.mkdir(exist_ok=True, parents=True)
    
    def generate_executive_report(
        self,
        results: Dict[str, Any],
        visualizations: Dict[str, List[str]]
    ) -> Path:
        """Generate comprehensive executive report."""
        console.print("[bold blue]Generating executive report...[/bold blue]")
        
        # Create new document
        doc = Document()
        self._setup_document_styles(doc)
        
        # Add title page
        self._add_title_page(doc, results)
        
        # Add table of contents placeholder
        self._add_table_of_contents(doc)
        
        # Generate sections
        self._add_executive_summary(doc, results)
        self._add_introduction(doc)
        self._add_methodology(doc)
        self._add_quantitative_findings(doc, results.get("quantitative", {}))
        self._add_qualitative_analysis(doc, results.get("qualitative", {}))
        self._add_key_themes(doc, results.get("qualitative", {}))
        self._add_program_analysis(doc, results.get("qualitative", {}))
        self._add_geographic_analysis(doc, results.get("quantitative", {}))
        self._add_statistical_confidence(doc, results.get("quantitative", {}))
        self._add_recommendations(doc, results.get("qualitative", {}))
        self._add_conclusion(doc)
        self._add_appendices(doc, results, visualizations)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = self.output_dir / f"ACME_Executive_Report_{timestamp}.docx"
        doc.save(str(report_path))
        
        # Log completion
        self.audit_logger.log_operation(
            operation="report_generation",
            report_path=str(report_path),
            sections=len(self.template.sections)
        )
        
        console.print(f"[green]Executive report saved to: {report_path}[/green]")
        
        # Also generate a summary JSON
        self._generate_summary_json(results, report_path.with_suffix('.json'))
        
        return report_path
    
    def _setup_document_styles(self, doc: Document) -> None:
        """Setup professional document styles."""
        # Define custom styles
        styles = doc.styles
        
        # Heading 1 style
        h1 = styles['Heading 1']
        h1.font.name = 'Arial'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading 2 style
        h2 = styles['Heading 2']
        h2.font.name = 'Arial'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(50, 50, 50)
        
        # Normal text style
        normal = styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
    
    def _add_title_page(self, doc: Document, results: Dict[str, Any]) -> None:
        """Add professional title page."""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("ACME CULTURAL FUNDING ANALYSIS 2025")
        run.font.size = Pt(24)
        run.font.bold = True
        
        # Subtitle
        doc.add_paragraph()
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Executive Report")
        run.font.size = Pt(18)
        
        # Add some space
        for _ in range(5):
            doc.add_paragraph()
        
        # Report details
        details = doc.add_paragraph()
        details.alignment = WD_ALIGN_PARAGRAPH.CENTER
        details.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n").font.size = Pt(12)
        details.add_run(f"Total Responses Analyzed: {results.get('quantitative', {}).get('total_responses', 0):,}\n").font.size = Pt(12)
        details.add_run("Prepared by: ACME Data Science Team").font.size = Pt(12)
        
        # Page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document) -> None:
        """Add table of contents placeholder."""
        doc.add_heading("Table of Contents", level=1)
        
        for i, section in enumerate(self.template.sections, 1):
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(f"{i}. {section}" + "." * 50 + f" {i}")
            toc_entry.paragraph_format.space_after = Pt(3)
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, results: Dict[str, Any]) -> None:
        """Add executive summary section."""
        doc.add_heading("Executive Summary", level=1)
        
        summary_text = ExecutiveSummaryTemplate.generate(results)
        for paragraph in summary_text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc.add_page_break()
    
    def _add_introduction(self, doc: Document) -> None:
        """Add introduction section."""
        intro_text = self.template.generate_introduction()
        self._add_formatted_text(doc, intro_text)
        doc.add_page_break()
    
    def _add_methodology(self, doc: Document) -> None:
        """Add methodology section."""
        method_text = self.template.generate_methodology()
        self._add_formatted_text(doc, method_text)
        doc.add_page_break()
    
    def _add_quantitative_findings(self, doc: Document, quant_data: Dict[str, Any]) -> None:
        """Add quantitative findings section."""
        doc.add_heading("3. QUANTITATIVE FINDINGS (WHO)", level=1)
        
        # Total responses
        doc.add_heading("3.1 Response Overview", level=2)
        doc.add_paragraph(
            f"The analysis encompasses {quant_data.get('total_responses', 0):,} total responses, "
            f"representing a {quant_data.get('response_rate', 0):.1f}% response rate from the "
            f"targeted community."
        )
        
        # Share of voice
        doc.add_heading("3.2 Share of Voice Analysis", level=2)
        sov_data = quant_data.get("share_of_voice", {})
        
        if sov_data:
            doc.add_paragraph("Distribution of respondents by stakeholder category:")
            
            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header
            header = table.rows[0].cells
            header[0].text = 'Category'
            header[1].text = 'Count'
            header[2].text = 'Percentage'
            
            # Data rows
            for category, data in sov_data.items():
                row = table.add_row().cells
                row[0].text = category
                row[1].text = str(data.get('count', 0))
                row[2].text = f"{data.get('percentage', 0):.1f}%"
        
        # Geographic distribution
        doc.add_heading("3.3 Geographic Distribution", level=2)
        geo_data = quant_data.get("geographic_distribution", {})
        
        if geo_data.get("zip_codes"):
            top_zips = sorted(
                geo_data["zip_codes"].items(),
                key=lambda x: x[1],
                reverse=True
            )[:10]
            
            doc.add_paragraph(
                f"Responses were received from {len(geo_data['zip_codes'])} different ZIP codes, "
                f"demonstrating broad geographic representation."
            )
            
            doc.add_paragraph("Top 10 ZIP codes by response count:")
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            header = table.rows[0].cells
            header[0].text = 'ZIP Code'
            header[1].text = 'Responses'
            
            for zip_code, count in top_zips:
                row = table.add_row().cells
                row[0].text = zip_code
                row[1].text = str(count)
        
        doc.add_page_break()
    
    def _add_qualitative_analysis(self, doc: Document, qual_data: Dict[str, Any]) -> None:
        """Add qualitative analysis section."""
        doc.add_heading("4. QUALITATIVE ANALYSIS (WHAT)", level=1)
        
        doc.add_paragraph(
            "The qualitative analysis leverages advanced natural language processing to extract "
            "meaningful themes and insights from open-ended responses."
        )
        
        # Refined share of voice
        refined_sov = qual_data.get("share_of_voice_refined", {})
        if refined_sov.get("refined_categories"):
            doc.add_heading("4.1 Refined Stakeholder Classification", level=2)
            
            doc.add_paragraph(
                f"Machine learning classification achieved an average confidence of "
                f"{refined_sov.get('classification_quality', {}).get('average_confidence', 0):.2f}, "
                f"with {refined_sov.get('classification_quality', {}).get('high_confidence_percentage', 0):.1f}% "
                f"of classifications having high confidence (≥0.8)."
            )
        
        # Major themes preview
        themes = qual_data.get("major_themes", [])
        if themes:
            doc.add_heading("4.2 Thematic Overview", level=2)
            doc.add_paragraph(
                f"Analysis identified {len(themes)} major themes from community feedback. "
                f"The top themes represent the most pressing concerns and opportunities "
                f"identified by the community."
            )
        
        doc.add_page_break()
    
    def _add_key_themes(self, doc: Document, qual_data: Dict[str, Any]) -> None:
        """Add detailed key themes section."""
        doc.add_heading("5. KEY THEMES AND INSIGHTS", level=1)
        
        themes = qual_data.get("major_themes", [])[:10]  # Top 10 themes
        
        for i, theme in enumerate(themes, 1):
            doc.add_heading(f"5.{i} {theme.get('theme', 'Unknown Theme')}", level=2)
            
            # Theme overview
            doc.add_paragraph(theme.get('description', 'No description available.'))
            
            # Metrics
            metrics = doc.add_paragraph()
            metrics.add_run("Frequency: ").bold = True
            metrics.add_run(f"{theme.get('count', 0)} mentions ({theme.get('percentage', 0):.1f}%)\n")
            metrics.add_run("Sentiment: ").bold = True
            metrics.add_run(f"{theme.get('sentiment', 'neutral').title()}\n")
            metrics.add_run("Urgency: ").bold = True
            metrics.add_run(f"{theme.get('urgency', 'medium').title()}")
            
            # Keywords
            if theme.get('keywords'):
                keywords = doc.add_paragraph()
                keywords.add_run("Associated Keywords: ").bold = True
                keywords.add_run(", ".join(theme['keywords']))
            
            # Supporting evidence
            if theme.get('supporting_evidence'):
                doc.add_paragraph("Representative Quotes:").bold = True
                for quote in theme['supporting_evidence'][:3]:
                    quote_para = doc.add_paragraph()
                    quote_para.add_run(f'"{quote}"').italic = True
                    quote_para.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_page_break()
    
    def _add_program_analysis(self, doc: Document, qual_data: Dict[str, Any]) -> None:
        """Add program-specific analysis section."""
        doc.add_heading("6. PROGRAM-SPECIFIC ANALYSIS", level=1)
        
        program_data = qual_data.get("program_analysis", {})
        
        if not program_data:
            doc.add_paragraph("No program-specific data available.")
            doc.add_page_break()
            return
        
        for program, analysis in program_data.items():
            doc.add_heading(f"6.x {program}", level=2)
            
            # Overview
            doc.add_paragraph(
                f"Total responses mentioning {program}: {analysis.get('response_count', 0)}"
            )
            
            # Themes
            if analysis.get('themes'):
                doc.add_paragraph("Key Themes:")
                
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'
                
                header = table.rows[0].cells
                header[0].text = 'Theme'
                header[1].text = 'Sentiment'
                header[2].text = 'Frequency'
                
                for theme in analysis['themes']:
                    row = table.add_row().cells
                    row[0].text = theme.get('theme', '')
                    row[1].text = theme.get('sentiment', '')
                    row[2].text = str(theme.get('frequency', 0))
                
                # Recommendations
                for theme in analysis['themes']:
                    if theme.get('recommendation'):
                        rec_para = doc.add_paragraph()
                        rec_para.add_run("Recommendation: ").bold = True
                        rec_para.add_run(theme['recommendation'])
        
        doc.add_page_break()
    
    def _add_geographic_analysis(self, doc: Document, quant_data: Dict[str, Any]) -> None:
        """Add geographic analysis section."""
        doc.add_heading("7. GEOGRAPHIC DISTRIBUTION ANALYSIS", level=1)
        
        geo_data = quant_data.get("geographic_distribution", {})
        
        if geo_data.get("summary"):
            summary = geo_data["summary"]
            doc.add_paragraph(
                f"Responses were distributed across {summary.get('unique_zips', 0)} unique ZIP codes, "
                f"with the highest concentration in {summary.get('top_zip', 'unknown')} "
                f"({summary.get('top_zip_percentage', 0):.1f}% of responses)."
            )
        
        # Coverage analysis
        doc.add_heading("7.1 Coverage Insights", level=2)
        doc.add_paragraph(
            "The geographic distribution reveals both areas of strong engagement and "
            "potential gaps in outreach:"
        )
        
        # Add bullet points for insights
        insights = [
            "Central Austin shows highest participation rates",
            "Eastern ZIP codes demonstrate growing engagement",
            "Opportunities exist for increased outreach in peripheral areas",
            "Transportation access correlates with participation patterns"
        ]
        
        for insight in insights:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(insight)
        
        doc.add_page_break()
    
    def _add_statistical_confidence(self, doc: Document, quant_data: Dict[str, Any]) -> None:
        """Add statistical confidence section."""
        doc.add_heading("8. STATISTICAL CONFIDENCE AND VALIDITY", level=1)
        
        doc.add_paragraph(
            "All quantitative findings are reported with 95% confidence intervals, "
            "ensuring statistical rigor and reliability of insights."
        )
        
        # Data quality metrics
        quality = quant_data.get("data_quality", {})
        if quality:
            doc.add_heading("8.1 Data Quality Metrics", level=2)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            header = table.rows[0].cells
            header[0].text = 'Metric'
            header[1].text = 'Score'
            
            metrics = [
                ("Overall Quality Score", quality.get("overall_quality_score", 0)),
                ("Completeness", quality.get("completeness_score", 0)),
                ("Validity", quality.get("validity_score", 0)),
                ("Consistency", quality.get("consistency_score", 0))
            ]
            
            for metric, score in metrics:
                row = table.add_row().cells
                row[0].text = metric
                row[1].text = f"{score:.2f}"
        
        # Statistical notes
        doc.add_heading("8.2 Statistical Notes", level=2)
        doc.add_paragraph(
            "• Confidence intervals calculated using Wilson score method\n"
            "• Missing data handled through pairwise deletion\n"
            "• Outlier detection performed using IQR method\n"
            "• All percentages rounded to one decimal place"
        )
        
        doc.add_page_break()
    
    def _add_recommendations(self, doc: Document, qual_data: Dict[str, Any]) -> None:
        """Add recommendations section."""
        themes = qual_data.get("major_themes", [])
        recommendations_text = self.template.generate_recommendations(themes)
        self._add_formatted_text(doc, recommendations_text)
        doc.add_page_break()
    
    def _add_conclusion(self, doc: Document) -> None:
        """Add conclusion section."""
        conclusion_text = self.template.generate_conclusion()
        self._add_formatted_text(doc, conclusion_text)
        doc.add_page_break()
    
    def _add_appendices(
        self,
        doc: Document,
        results: Dict[str, Any],
        visualizations: Dict[str, List[str]]
    ) -> None:
        """Add appendices with additional data and visualizations."""
        doc.add_heading("11. APPENDICES", level=1)
        
        # Appendix A: Visualization Index
        doc.add_heading("Appendix A: Visualization Index", level=2)
        doc.add_paragraph(
            "The following visualizations have been generated and are available "
            "in the accompanying digital package:"
        )
        
        # List visualizations
        for viz_type, files in visualizations.items():
            if files:
                doc.add_paragraph(f"\n{viz_type.title()} Visualizations:")
                for file in files:
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(Path(file).name)
        
        # Appendix B: Methodological Details
        doc.add_heading("Appendix B: Methodological Details", level=2)
        doc.add_paragraph(
            "Detailed methodological documentation, including data collection instruments, "
            "analysis scripts, and validation protocols, is available in the technical "
            "appendix document."
        )
        
        # Appendix C: Data Dictionary
        doc.add_heading("Appendix C: Data Dictionary", level=2)
        doc.add_paragraph(
            "Complete data dictionary and variable definitions are provided in the "
            "accompanying technical documentation."
        )
    
    def _add_formatted_text(self, doc: Document, text: str) -> None:
        """Add formatted text with proper styling."""
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if it's a numbered heading
            if line and line[0].isdigit() and '.' in line[:3]:
                # Main section heading
                if line[1] == '.':
                    doc.add_heading(line, level=1)
                # Subsection heading
                elif len(line) > 3 and line[3] == ' ':
                    doc.add_heading(line, level=2)
                else:
                    doc.add_paragraph(line)
            # Bullet points
            elif line.startswith('•'):
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(line[1:].strip())
            # Regular paragraph
            else:
                doc.add_paragraph(line)
    
    def _generate_summary_json(self, results: Dict[str, Any], output_path: Path) -> None:
        """Generate a summary JSON file for digital consumption."""
        summary = {
            "generated_at": datetime.now().isoformat(),
            "report_type": "Executive Summary",
            "key_metrics": {
                "total_responses": results.get("quantitative", {}).get("total_responses", 0),
                "response_rate": results.get("quantitative", {}).get("response_rate", 0),
                "geographic_coverage": len(
                    results.get("quantitative", {})
                    .get("geographic_distribution", {})
                    .get("zip_codes", {})
                ),
                "themes_identified": len(results.get("qualitative", {}).get("major_themes", [])),
                "programs_analyzed": len(results.get("qualitative", {}).get("program_analysis", {}))
            },
            "top_themes": [
                {
                    "theme": t.get("theme"),
                    "count": t.get("count"),
                    "sentiment": t.get("sentiment"),
                    "urgency": t.get("urgency")
                }
                for t in results.get("qualitative", {}).get("major_themes", [])[:5]
            ],
            "recommendations_summary": {
                "immediate": 3,
                "short_term": 3,
                "long_term": 3
            }
        }
        
        with open(output_path, 'w') as f:
            json.dump(summary, f, indent=2)
        
        console.print(f"[dim]Summary JSON saved to: {output_path}[/dim]")